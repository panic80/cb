"""Loaders for Office documents (DOCX, XLSX, CSV)."""

import asyncio
import csv
import io
from typing import List, Dict, Any, Optional
from datetime import datetime
import pandas as pd
from docx import Document as DocxDocument
import openpyxl

from langchain_core.documents import Document

from app.core.logging import get_logger
from app.models.documents import DocumentType

logger = get_logger(__name__)


class DocxLoader:
    """Loader for Microsoft Word documents."""
    
    def __init__(self, file_path: str):
        """Initialize DOCX loader."""
        self.file_path = file_path
        
    async def load(self) -> List[Document]:
        """Load and parse DOCX content."""
        try:
            # Run in executor since it's sync
            loop = asyncio.get_event_loop()
            documents = await loop.run_in_executor(None, self._load_sync)
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load DOCX from {self.file_path}: {e}")
            raise
            
    def _load_sync(self) -> List[Document]:
        """Synchronous DOCX loading."""
        doc = DocxDocument(self.file_path)
        
        # Extract metadata
        metadata = self._extract_metadata(doc)
        
        # Extract content with structure preservation
        content_parts = []
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    level = para.style.name[-1] if para.style.name[-1].isdigit() else '1'
                    content_parts.append(f"\n\n{'#' * int(level)} {para.text}\n")
                else:
                    content_parts.append(para.text)
                    
        # Process tables
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                content_parts.append(f"\n{table_text}\n")
                
        # Join content
        content = "\n".join(content_parts)
        
        # Create document
        return [Document(
            page_content=content,
            metadata={
                "source": self.file_path,
                "type": DocumentType.DOCX,
                **metadata
            }
        )]
        
    def _extract_metadata(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract metadata from DOCX."""
        metadata = {}
        
        try:
            # Extract core properties
            core_props = doc.core_properties
            
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["last_modified"] = core_props.modified.isoformat()
                
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
            
        return metadata
        
    def _extract_table(self, table) -> str:
        """Extract table content as formatted text."""
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(" | ".join(cells))
            
        return "\n".join(rows)


class XlsxLoader:
    """Loader for Excel spreadsheets."""
    
    def __init__(self, file_path: str):
        """Initialize XLSX loader."""
        self.file_path = file_path
        
    async def load(self) -> List[Document]:
        """Load and parse XLSX content."""
        try:
            # Run in executor since it's sync
            loop = asyncio.get_event_loop()
            documents = await loop.run_in_executor(None, self._load_sync)
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load XLSX from {self.file_path}: {e}")
            raise
            
    def _load_sync(self) -> List[Document]:
        """Synchronous XLSX loading."""
        documents = []
        
        # Load workbook
        wb = openpyxl.load_workbook(self.file_path, read_only=True, data_only=True)
        
        # Process each sheet
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            
            # Skip empty sheets
            if sheet.max_row == 0:
                continue
                
            # Extract content
            content = self._extract_sheet_content(sheet)
            
            if content:
                # Create document for this sheet
                doc = Document(
                    page_content=content,
                    metadata={
                        "source": self.file_path,
                        "type": DocumentType.XLSX,
                        "sheet_name": sheet_name,
                        "sheet_index": wb.sheetnames.index(sheet_name),
                        "total_sheets": len(wb.sheetnames),
                        "loaded_at": datetime.utcnow().isoformat()
                    }
                )
                documents.append(doc)
                
        wb.close()
        return documents
        
    def _extract_sheet_content(self, sheet) -> str:
        """Extract content from a sheet."""
        content_parts = []
        
        # Try to use pandas for better table extraction
        try:
            # Convert to pandas DataFrame
            data = sheet.values
            df = pd.DataFrame(data)
            
            # Use first row as headers if it looks like headers
            if df.shape[0] > 1:
                first_row = df.iloc[0]
                if all(isinstance(val, str) for val in first_row if val is not None):
                    df.columns = first_row
                    df = df[1:]
                    
            # Convert to markdown table
            content = df.to_markdown(index=False)
            content_parts.append(content)
            
        except Exception as e:
            logger.warning(f"Failed to use pandas for sheet extraction: {e}")
            
            # Fallback to manual extraction
            rows = []
            for row in sheet.iter_rows(values_only=True):
                # Filter out completely empty rows
                if any(cell is not None for cell in row):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    rows.append(row_text)
                    
            content_parts.append("\n".join(rows))
            
        return "\n".join(content_parts)


class CsvLoader:
    """Loader for CSV files."""
    
    def __init__(self, file_path: str, encoding: str = 'utf-8'):
        """Initialize CSV loader."""
        self.file_path = file_path
        self.encoding = encoding
        
    async def load(self) -> List[Document]:
        """Load and parse CSV content."""
        try:
            # Run in executor since it's sync
            loop = asyncio.get_event_loop()
            documents = await loop.run_in_executor(None, self._load_sync)
            return documents
            
        except Exception as e:
            logger.error(f"Failed to load CSV from {self.file_path}: {e}")
            raise
            
    def _load_sync(self) -> List[Document]:
        """Synchronous CSV loading."""
        try:
            # Try pandas first for better handling
            df = pd.read_csv(self.file_path, encoding=self.encoding)
            
            # Convert to markdown table
            content = df.to_markdown(index=False)
            
            # Extract metadata
            metadata = {
                "source": self.file_path,
                "type": DocumentType.CSV,
                "encoding": self.encoding,
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "loaded_at": datetime.utcnow().isoformat()
            }
            
            return [Document(
                page_content=content,
                metadata=metadata
            )]
            
        except Exception as e:
            logger.warning(f"Failed to use pandas for CSV: {e}")
            
            # Fallback to standard CSV reader
            content_parts = []
            row_count = 0
            
            with open(self.file_path, 'r', encoding=self.encoding) as f:
                reader = csv.reader(f)
                
                for i, row in enumerate(reader):
                    if i == 0:
                        # Assume first row is headers
                        headers = row
                        content_parts.append(" | ".join(headers))
                        content_parts.append(" | ".join(["-" * len(h) for h in headers]))
                    else:
                        content_parts.append(" | ".join(row))
                        row_count += 1
                        
            content = "\n".join(content_parts)
            
            return [Document(
                page_content=content,
                metadata={
                    "source": self.file_path,
                    "type": DocumentType.CSV,
                    "encoding": self.encoding,
                    "rows": row_count,
                    "loaded_at": datetime.utcnow().isoformat()
                }
            )]